import streamlit as st
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
from PIL import Image
import json
import io

st.set_page_config(layout="wide")

# Funktion zum Erstellen eines Donut-Charts
def create_donut_chart(percentage):
    if percentage <= 25:
        color = 'red'
    elif percentage <= 50:
        color = 'orange'
    elif percentage <= 75:
        color = '#1f77b4'
    else:
        color = 'green'
    
    fig, ax = plt.subplots()
    ax.pie([percentage, 100 - percentage], 
           labels=['', ''], 
           startangle=90, 
           colors=[color, '#d3d3d3'], 
           wedgeprops={'width': 0.3})
    
    # Prozentsatz in der Mitte des Donuts anzeigen
    ax.text(0, 0, f"{percentage}%", ha='center', va='center', fontsize=20, color=color)
    
    # Chart als Bild in einem BytesIO-Objekt speichern
    buf = BytesIO()
    plt.savefig(buf, format="png")
    buf.seek(0)
    plt.close(fig)
    return buf

# Streamlit App
st.title("Brandschutzgutachten - Tabelle mit Fertigstellungsgrad")

# Daten basierend auf den Brandschutzanforderungen
data = [
    {"anforderung": "Brandschutztüren in Flucht- und Rettungswegen", 
     "maßnahme": "Installation von feuerhemmenden Türen (T30)", 
     "wirtschaftliche_maßnahme": "Einsatz von zugelassenen T30-Brandschutztüren aus Stahl", 
     "percentage": 100},
    
    {"anforderung": "Brandmeldeanlagen (DIN 14675)", 
     "maßnahme": "Zentrale Brandmeldeanlage (BMA)", 
     "wirtschaftliche_maßnahme": "Rauchmelder nur in den Wohnungen", 
     "percentage": 70},
    
    {"anforderung": "Rauchabzugsanlagen in Treppenhäusern", 
     "maßnahme": "Mechanische Rauch- und Wärmeabzugsanlagen", 
     "wirtschaftliche_maßnahme": "Natürliche Rauchabzugsanlagen (Fenster)", 
     "percentage": 85},
    
    {"anforderung": "Flucht- und Rettungswege", 
     "maßnahme": "Freihaltung und Kennzeichnung durch Sicherheitsbeleuchtung", 
     "wirtschaftliche_maßnahme": "Einfache Beleuchtung ohne Sicherheitsbeleuchtung", 
     "percentage": 60},
    
    {"anforderung": "Feuerwiderstandsfähige Decken und Wände (DIN 4102)", 
     "maßnahme": "Feuerbeständige Materialien (F90)", 
     "wirtschaftliche_maßnahme": "Feuerhemmende Materialien (F30)", 
     "percentage": 75},
    
    {"anforderung": "Sicherheitsbeleuchtung (DIN EN 1838)", 
     "maßnahme": "Sicherheitsbeleuchtung an Notausgängen", 
     "wirtschaftliche_maßnahme": "Notbeleuchtung nur an Treppenabgängen", 
     "percentage": 50},
    
    {"anforderung": "Löschwasserversorgung (DIN 1988-600)", 
     "maßnahme": "Vorhaltung einer Löschwasseranlage", 
     "wirtschaftliche_maßnahme": "Anschluss an externes Löschwassernetz", 
     "percentage": 80},
    
    {"anforderung": "Abschottung von Leitungen (DIN 4102-11)", 
     "maßnahme": "Feuerwiderstandsfähige Abschottung von Leitungen", 
     "wirtschaftliche_maßnahme": "Teilweise Abschottung bei Hauptleitungen", 
     "percentage": 70},
    
    {"anforderung": "Aufstellflächen für Feuerwehrfahrzeuge", 
     "maßnahme": "Bereitstellung geeigneter Aufstellflächen", 
     "wirtschaftliche_maßnahme": "Minimalistische Auslegung der Flächen", 
     "percentage": 90},
    
    {"anforderung": "Notrufeinrichtungen in Aufzügen (DIN EN 81-28)", 
     "maßnahme": "Permanente Überwachung der Notrufsysteme", 
     "wirtschaftliche_maßnahme": "Einfache Notrufsysteme ohne permanente Überwachung", 
     "percentage": 70},
]

# Tabelle mit vier Spalten anzeigen und bearbeiten
st.write("### Anforderungen, Maßnahmen, wirtschaftliche Umsetzung und Erfüllungsgrad")

# Überschriften
st.write(
    """
    | **Anforderung** | **Erforderliche Maßnahme** | **Wirtschaftliche Maßnahme** | **Erfüllungsgrad** |
    |-----------------|----------------------------|------------------------------|--------------------|
    """
)

# Inhalte
for row in data:
    col1, col2, col3, col4 = st.columns([2, 2, 2, 1])

    # Inhalte der ersten drei Spalten anzeigen
    col1.write(row["anforderung"])
    col2.write(row["maßnahme"])
    col3.write(row["wirtschaftliche_maßnahme"])

    # Slider in der vierten Spalte anzeigen
    row["percentage"] = col3.slider(f"Erfüllungsgrad", min_value=0, max_value=100, value=row["percentage"], step=5)

    # Donut-Chart als Indikator anzeigen
    donut_chart = create_donut_chart(row["percentage"])
    col4.image(donut_chart, use_column_width=True)

# Speichern Button
if st.button('Speichern'):
    # Word-Dokument erstellen
    doc = Document()
    doc.add_heading('Brandschutzgutachten - Tabelle mit Fertigstellungsgrad', 0)
    
    # Tabelle im Word-Dokument erstellen
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Anforderung'
    hdr_cells[1].text = 'Erforderliche Maßnahme'
    hdr_cells[2].text = 'Wirtschaftliche Maßnahme'
    hdr_cells[3].text = 'Fertigstellungsgrad'
    
    for row in data:
        # Neue Zeile in der Tabelle erstellen
        row_cells = table.add_row().cells
        
        # Text in den ersten drei Spalten
        row_cells[0].text = row["anforderung"]
        row_cells[1].text = row["maßnahme"]
        row_cells[2].text = row["wirtschaftliche_maßnahme"]
        
        # Donut-Chart als Bild hinzufügen
        donut_chart = create_donut_chart(row["percentage"])
        
        # Bild in der letzten Spalte speichern und hinzufügen
        image_stream = Image.open(donut_chart)
        image_path = f"temp_{row['percentage']}.png"
        image_stream.save(image_path)
        row_cells[3].add_paragraph().add_run().add_picture(image_path, width=Inches(1.5))

    # Word-Dokument speichern
    doc.save('Brandschutzgutachten.docx')
    st.success("Das Brandschutzgutachten wurde erfolgreich gespeichert!")


    # Word-Dokument speichern
    doc.save('Brandschutzgutachten.docx')
    st.success("Das Brandschutzgutachten wurde erfolgreich gespeichert!")

    # Speichern des Word-Dokuments
    #doc.save('Tabelle_Fertigstellungsgrad.docx')
    #st.success("Das Word-Dokument wurde erfolgreich erstellt und gespeichert.")

     # Das Word-Dokument in einem BytesIO-Objekt speichern
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    # Download Button anzeigen
    st.download_button(
        label="Download Word-Dokument",
        data=doc_io,
        file_name="Tabelle_Fertigstellungsgrad.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# Beispielhafter SessionState
if 'my_state' not in st.session_state:
    st.session_state.my_state = {'key1': 'value1', 'key2': 42}

# Funktion zum Speichern des SessionState in JSON und Download anbieten
def save_sessionstate_to_json():
    # SessionState in ein JSON-kompatibles Format umwandeln
    session_dict = {k: v for k, v in st.session_state.items()}
    
    # JSON in einen BytesIO-Stream schreiben
    json_str = json.dumps(session_dict, indent=4)
    json_bytes = json_str.encode('utf-8')
    json_io = io.BytesIO(json_bytes)
    
    # Download-Button anzeigen
    st.sidebar.download_button(
        label="Aktuelles Projekt speichern",
        data=json_io,
        file_name="session_state.json",
        mime="application/json"
    )


# Download-Button anzeigen
save_sessionstate_to_json()

def upload_sessionstate_from_json(uploaded_file):
    if uploaded_file is not None:
        # JSON-Daten lesen und in den SessionState schreiben
        session_dict = json.load(uploaded_file)
        st.session_state.update(session_dict)
        st.success("SessionState erfolgreich aktualisiert!")


# JSON-Datei hochladen
uploaded_file = st.sidebar.file_uploader("Bestehendes Projekt per JSON Datei hochladen", type="json")

# Button zum Hochladen und SessionState aktualisieren
if uploaded_file is not None:
    upload_sessionstate_from_json(uploaded_file)
