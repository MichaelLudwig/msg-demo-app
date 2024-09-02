import streamlit as st
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
from PIL import Image

st.page_link("streamlit_app.py", label="Home", icon="🏠")
st.page_link("pages/charts.py", label="Chart Demo", icon="1️⃣")

# Funktion zum Erstellen eines Donut-Charts
def create_donut_chart(percentage):
    if percentage <= 25:
        color = 'red'
    elif percentage <= 50:
        color = 'orange'
    elif percentage <= 75:
        color = '#1f77b4'
    else:
        color = 'green'
    
    fig, ax = plt.subplots()
    ax.pie([percentage, 100 - percentage], 
           labels=['', ''], 
           startangle=90, 
           colors=[color, '#d3d3d3'], 
           wedgeprops={'width': 0.3})
    
    # Prozentsatz in der Mitte des Donuts anzeigen
    ax.text(0, 0, f"{percentage}%", ha='center', va='center', fontsize=20, color=color)
    
    # Chart als Bild in einem BytesIO-Objekt speichern
    buf = BytesIO()
    plt.savefig(buf, format="png")
    buf.seek(0)
    plt.close(fig)
    return buf

# Streamlit App
st.title("Tabelle mit Fertigstellungsgrad")

# Dummy-Daten
data = [
    {"text": "Lorem ipsum dolor sit amet, consectetur adipiscing. Sed do eiusmod tempor incididunt ut labore et dolore.", "percentage": 15},
    {"text": "Sed do eiusmod tempor incididunt ut labore et dolore. Magna aliqua. Ut enim ad minim veniam, quis nostrud.", "percentage": 30},
    {"text": "Magna aliqua. Ut enim ad minim veniam, quis nostrud. Exercitation ullamco laboris nisi ut aliquip ex ea.", "percentage": 55},
    {"text": "Exercitation ullamco laboris nisi ut aliquip ex ea. Lorem ipsum dolor sit amet, consectetur adipiscing. ", "percentage": 90},
]

# Tabelle anzeigen und bearbeiten
for row in data:
    col1, col2 = st.columns([3, 1])

    # Text in der ersten Spalte anzeigen
    col1.write(row["text"])

    
    # Slider zur Anpassung des Prozentsatzes
    row["percentage"] = col1.slider(f"Prozentwert für diesen Eintrag:", min_value=0, max_value=100, value=row["percentage"], step=5)
    
    # Schließen des farbigen Containers
    st.markdown("</div>", unsafe_allow_html=True)
    
    # Donut-Chart in der zweiten Spalte anzeigen
    donut_chart = create_donut_chart(row["percentage"])
    col2.image(donut_chart, use_column_width=True)

# Speichern Button
if st.button('Speichern'):
    # Word-Dokument erstellen
    doc = Document()
    doc.add_heading('Tabelle mit Fertigstellungsgrad', 0)
    
    # Tabelle im Word-Dokument erstellen
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Beschreibung'
    hdr_cells[1].text = 'Fertigstellungsgrad'
    
    for row in data:
        # Neue Zeile in der Tabelle erstellen
        row_cells = table.add_row().cells
        
        # Text in der ersten Spalte
        row_cells[0].text = row["text"]
        
        # Donut-Chart in der zweiten Spalte
        donut_chart = create_donut_chart(row["percentage"])
        
        # Bild in der zweiten Spalte speichern und hinzufügen
        image_stream = Image.open(donut_chart)
        image_path = f"temp_{row['percentage']}.png"
        image_stream.save(image_path)
        row_cells[1].add_paragraph().add_run().add_picture(image_path, width=Inches(1.5))
    
    # Speichern des Word-Dokuments
    #doc.save('Tabelle_Fertigstellungsgrad.docx')
    #st.success("Das Word-Dokument wurde erfolgreich erstellt und gespeichert.")

     # Das Word-Dokument in einem BytesIO-Objekt speichern
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    # Download Button anzeigen
    st.download_button(
        label="Download Word-Dokument",
        data=doc_io,
        file_name="Tabelle_Fertigstellungsgrad.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
